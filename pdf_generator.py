"""
pdf_generator.py - FIX DEFINITIVO - scorpio/escorpio + sin azteca
Usa Portadas/ y Contenido/
"""

import os
import random
import logging
import unicodedata
from datetime import datetime
from docx import Document

from pdf_co_integration import merge_docx_to_pdf, docx_to_pdf

logger = logging.getLogger(__name__)

COVERS_DIR = os.path.join(os.path.dirname(__file__), "Portadas")
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "Contenido")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "generated_pdfs")

os.makedirs(OUTPUT_DIR, exist_ok=True)

CONTENT_TEMPLATE_ORDER = [
    "contenido_signo_occidental.docx",
    "contenido_signo_chino.docx",
    "contenido_signo_celta.docx",
    "contenido_signo_maya.docx",
    "contenido_signo_egipcio.docx",
    "cierre.docx",
]

# Mapa para que escorpio = scorpio, tauro = taurus, etc
SIGNOS_MAP = {
    "aries": ["aries"],
    "tauro": ["tauro", "taurus"],
    "taurus": ["tauro", "taurus"],
    "geminis": ["geminis", "gemini"],
    "gemini": ["geminis", "gemini"],
    "cancer": ["cancer"],
    "leo": ["leo"],
    "virgo": ["virgo"],
    "libra": ["libra"],
    "escorpio": ["escorpio", "scorpio", "escorpion"],
    "scorpio": ["escorpio", "scorpio", "escorpion"],
    "escorpion": ["escorpio", "scorpio"],
    "sagitario": ["sagitario", "sagittarius"],
    "sagittarius": ["sagitario"],
    "capricornio": ["capricornio", "capricorn"],
    "capricorn": ["capricornio", "capricorn"],
    "acuario": ["acuario", "aquarius"],
    "aquarius": ["acuario", "aquarius"],
    "piscis": ["piscis", "pisces"],
    "pisces": ["piscis", "pisces"],
}

def normalize_text(s: str) -> str:
    if not s:
        return ""
    nfkd = unicodedata.normalize('NFKD', s.lower())
    return "".join([c for c in nfkd if not unicodedata.combining(c)])

def _pick_random_cover(signo: str = "") -> str:
    if not os.path.isdir(COVERS_DIR):
        raise FileNotFoundError(f"No existe la carpeta de portadas: {COVERS_DIR}")

    all_files = [f for f in os.listdir(COVERS_DIR) if f.lower().endswith((".jpeg", ".jpg", ".png", ".pdf", ".docx"))]
    if not all_files:
        raise FileNotFoundError(f"No hay portadas en {COVERS_DIR}.")

    if signo:
        signo_norm = normalize_text(signo)
        variantes = SIGNOS_MAP.get(signo_norm, [signo_norm])
        filtered = []
        for f in all_files:
            f_norm = normalize_text(f)
            if any(v in f_norm for v in variantes):
                filtered.append(f)
        
        if filtered:
            pool = filtered[:3] if len(filtered) >= 3 else filtered
            chosen = random.choice(pool)
            logger.info(f"Portada elegida para {signo} ({signo_norm}) -> {chosen} de {filtered}")
            return chosen
        else:
            logger.warning(f"No hay portada para signo {signo} ({signo_norm}) variantes {variantes}. Archivos: {all_files[:15]} -> random")

    return random.choice(all_files)

def _fill_docx_template(template_path: str, output_path: str, context: dict):
    doc = Document(template_path)

    def replace_in_paragraph(paragraph):
        for key, value in context.items():
            placeholder = "{{" + key + "}}"
            if placeholder in paragraph.text:
                if isinstance(value, list):
                    if "tarot" in key:
                        val_str = "\n\n".join([f"Pregunta {i+1}: {r.get('question','')}\nCarta: {r.get('card','')} - Respuesta: {r.get('answer','')}\n{r.get('interpretation','')}" for i, r in enumerate(value)])
                    else:
                        val_str = "\n".join([str(v) for v in value])
                else:
                    val_str = str(value)
                
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, val_str)
                full_text = "".join(r.text for r in paragraph.runs)
                if placeholder in full_text:
                    new_text = full_text.replace(placeholder, val_str)
                    for i, run in enumerate(paragraph.runs):
                        run.text = new_text if i == 0 else ""

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    doc.save(output_path)

def build_natal_chart_pdf(full_name: str, birth_date: str, birth_time: str, reading: dict) -> tuple:
    cover_file = _pick_random_cover(reading.get("zodiac_western", ""))
    cover_path = os.path.join(COVERS_DIR, cover_file)

    timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    safe_name = "".join(c for c in full_name if c.isalnum() or c == " ").strip().replace(" ", "_")
    work_dir = os.path.join(OUTPUT_DIR, f"{safe_name}_{timestamp}")
    os.makedirs(work_dir, exist_ok=True)

    tarot_list = reading.get("tarot_readings", reading.get("tarot_lecturas", []))
    header_context = {
        "nombre": full_name,
        "fecha_nacimiento": birth_date,
        "hora_nacimiento": birth_time if birth_time else "No especificada",
        "lugar_nacimiento": reading.get("birth_place", "No especificado"),
        "genero": reading.get("birth_gender", "No especificado"),
        "signo_occidental": reading.get("zodiac_western", ""),
        "signo_chino": reading.get("zodiac_chinese", ""),
        "signo_celta": reading.get("zodiac_celtic", ""),
        "signo_maya": reading.get("zodiac_mayan", ""),
        "signo_egipcio": reading.get("zodiac_egyptian", ""),
        # azteca eliminado - no existe en DB
        "ascendente": reading.get("ascending_sign", ""),
        "luna": reading.get("moon_sign", ""),
        "interpretacion": reading.get("interpretation", reading.get("detailed_interpretation", "")),
        "interpretacion_detallada": reading.get("detailed_interpretation", reading.get("interpretation", "")),
        "tarot_lecturas": tarot_list,
        "tarot_readings": tarot_list,
        "mensaje_final": reading.get("overall_message", ""),
        "overall_message": reading.get("overall_message", ""),
        "fecha_lectura": datetime.utcnow().strftime("%d/%m/%Y"),
    }

    docx_files_in_order = []

    if cover_file.lower().endswith(".docx"):
        filled_cover = os.path.join(work_dir, "00_portada.docx")
        _fill_docx_template(cover_path, filled_cover, header_context)
        docx_files_in_order.append(filled_cover)
    else:
        docx_files_in_order.append(cover_path)

    for i, template_name in enumerate(CONTENT_TEMPLATE_ORDER, start=1):
        template_path = os.path.join(TEMPLATES_DIR, template_name)
        if not os.path.exists(template_path):
            logger.warning(f"Falta template en Contenido/: {template_name}, se omite.")
            continue
        filled_path = os.path.join(work_dir, f"{i:02d}_{template_name}")
        _fill_docx_template(template_path, filled_path, header_context)
        docx_files_in_order.append(filled_path)

    local_backup_path = os.path.join(work_dir, "carta_astral_final.pdf")
    final_pdf_url = merge_docx_to_pdf(docx_files_in_order, local_backup_path)

    logger.info(f"PDF generado: {final_pdf_url} con portada {cover_file} y {len(docx_files_in_order)} secciones")
    return final_pdf_url, cover_file

